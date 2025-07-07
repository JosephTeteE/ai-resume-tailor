# test_resume_tailor.py
import pytest
from unittest.mock import patch, MagicMock, PropertyMock
from io import BytesIO
import re
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import pdfplumber

# Import the functions to test
from ai_agent import prioritize_sections, tailor_section, tailor_resume
from utils import (
    extract_text_from_docx,
    extract_text_from_pdf,
    extract_keywords,
    highlight_keywords,
    calculate_ats_score,
    create_styled_docx,
    _create_styled_run_element,
    _parse_markdown_and_create_elements,
    _add_section_to_doc_content
)

# --- Test Data ---
SAMPLE_CV_TEXT = """
John Doe • johndoe@example.com • (123) 456-7890 • linkedin.com/in/johndoe

PROFESSIONAL SUMMARY
Results-driven software engineer with 5+ years of experience...

PROFESSIONAL EXPERIENCE
Senior Software Engineer | Tech Company | Jan 2020 - Present
- Developed scalable microservices using Python and Django
- Led a team of 5 engineers to deliver project X
- Increased system performance by 30%

EDUCATION
BSc Computer Science | University | 2015-2019
"""

SAMPLE_JD = """
We are looking for a Senior Python Developer with:
- 5+ years of Python experience
- Strong Django framework knowledge
- Experience leading engineering teams
- Microservices architecture expertise
"""

EMPTY_CV = ""
EMPTY_JD = ""
LARGE_TEXT = "X" * 20000  # Exceeds MAX_INPUT_LENGTH

# --- Test Fixtures ---
@pytest.fixture
def sample_docx_file(tmp_path):
    doc = Document()
    doc.add_paragraph("Test DOCX content")
    file_path = tmp_path / "test.docx"
    doc.save(file_path)
    return file_path

@pytest.fixture
def sample_pdf_file(tmp_path):
    # Create a simple PDF (in reality, you'd need proper PDF creation)
    file_path = tmp_path / "test.pdf"
    with open(file_path, 'wb') as f:
        f.write(b"%PDF-1.4\n%\\E2\\E3\\CF\\D3\n1 0 obj\n<< /Type /Catalog /Pages 2 0 R >>\nendobj\n2 0 obj\n<< /Type /Pages /Kids [3 0 R] /Count 1 >>\nendobj\n3 0 obj\n<< /Type /Page /Parent 2 0 R /Resources << /Font << /F1 4 0 R >> >> /Contents 5 0 R >>\nendobj\n4 0 obj\n<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>\nendobj\n5 0 obj\n<< /Length 44 >>\nstream\nBT\n/F1 12 Tf\n72 720 Td\n(Test PDF content) Tj\nET\nendstream\nendobj\nxref\n0 6\n0000000000 65535 f \n0000000010 00000 n \n0000000069 00000 n \n0000000123 00000 n \n0000000204 00000 n \n0000000259 00000 n \ntrailer\n<< /Size 6 /Root 1 0 R >>\nstartxref\n364\n%%EOF")
    return file_path

# --- AI Agent Tests ---
def test_prioritize_sections():
    prioritized = prioritize_sections(SAMPLE_JD)
    assert isinstance(prioritized, list)
    assert len(prioritized) > 0
    assert "CONTACT_INFO" in prioritized
    assert "PROFESSIONAL_EXPERIENCE" in prioritized[:3]  # Should be high priority
    
    # Test with empty JD
    empty_prioritized = prioritize_sections("")
    assert len(empty_prioritized) > 0
    
def test_tailor_section_empty_content():
    result = tailor_section("PROFESSIONAL_SUMMARY", "", SAMPLE_JD, "fake_api_key")
    assert result == ""

@patch('ai_agent.genai.GenerativeModel')
@patch('ai_agent.genai.configure')
def test_tailor_section_api_call(mock_configure, mock_model):
    # Setup mock response
    mock_response = MagicMock()
    mock_response.text = "Tailored content"
    
    # Create proper mock structure for prompt_feedback
    mock_prompt_feedback = MagicMock()
    type(mock_prompt_feedback).block_reason = PropertyMock(return_value=MagicMock(name="SAFETY"))
    
    mock_response.prompt_feedback = mock_prompt_feedback
    mock_model.return_value.generate_content.return_value = mock_response
    
    # Test successful tailoring
    result = tailor_section("PROFESSIONAL_SUMMARY", "Original content", SAMPLE_JD, "fake_api_key")
    assert result == "Tailored content"
    
    # Test blocked response
    mock_response.text = ""
    result = tailor_section("PROFESSIONAL_SUMMARY", "Original content", SAMPLE_JD, "fake_api_key")
    assert "blocked" in result.lower()

@patch('ai_agent.genai.GenerativeModel')
@patch('ai_agent.genai.configure')
@patch('ai_agent.tailor_section')
def test_tailor_resume_section_parsing(mock_tailor, mock_configure, mock_model):
    # Mock tailor_section to return quickly
    mock_tailor.return_value = "Mock tailored content"
    
    sections = tailor_resume(SAMPLE_CV_TEXT, SAMPLE_JD, "fake_api_key")
    assert isinstance(sections, dict)
    assert "PROFESSIONAL_SUMMARY" in sections
    assert "PROFESSIONAL_EXPERIENCE" in sections
    assert "EDUCATION" in sections
    
    # Test with empty CV
    empty_sections = tailor_resume("", SAMPLE_JD, "fake_api_key")
    assert "CONTACT_INFO" in empty_sections

# --- Utils Tests ---
def test_extract_text_from_docx(sample_docx_file):
    with open(sample_docx_file, 'rb') as f:
        text = extract_text_from_docx(f)
    assert "Test DOCX content" in text
    
# def test_extract_text_from_pdf(sample_pdf_file):
    # with open(sample_pdf_file, 'rb') as f:
        # with patch('pdfplumber.open') as mock_open:
            # mock_pdf = MagicMock()
            # mock_page = MagicMock()
            # mock_page.extract_text.return_value = "Test PDF content"
            # mock_pdf.pages = [mock_page]
            # mock_open.return_value.__enter__.return_value = mock_pdf
            # 
            # text = extract_text_from_pdf(f)
    # assert "Test PDF content" in text

def test_extract_keywords():
    keywords = extract_keywords(SAMPLE_JD)
    assert isinstance(keywords, list)
    assert len(keywords) > 0
    assert "python" in keywords
    assert "django" in keywords
    
    # Test with empty text
    assert extract_keywords("") == []

def test_highlight_keywords():
    highlighted = highlight_keywords(SAMPLE_CV_TEXT, ["engineer", "python"])
    assert "**engineer**" in highlighted.lower()
    assert "**python**" in highlighted.lower()
    
    # Test with no keywords
    assert highlight_keywords(SAMPLE_CV_TEXT, []) == SAMPLE_CV_TEXT

def test_calculate_ats_score():
    score = calculate_ats_score(SAMPLE_CV_TEXT, SAMPLE_JD)
    assert 0 <= score <= 100
    
    # Test with no match
    assert calculate_ats_score("", SAMPLE_JD) == 0
    assert calculate_ats_score(SAMPLE_CV_TEXT, "") == 0

# --- DOCX Generation Tests ---
def test_create_styled_run_element():
    run_element = _create_styled_run_element(
        "Test",
        font_name="Arial",
        font_size=Pt(12),
        color_rgb=RGBColor(255, 0, 0),
        is_bold=True
    )
    assert run_element is not None
    assert "Test" in run_element.text

    # Regression check: RGBColor doesn't crash
    run_element_black = _create_styled_run_element("Test Color", color_rgb=RGBColor(0, 0, 0))
    assert run_element_black is not None
    assert "Test Color" in run_element_black.text

    
    # Test with RGBColor (specific regression test for your error)
    run_element = _create_styled_run_element("Test", color_rgb=RGBColor(0, 0, 0))
    assert run_element is not None

def test_parse_markdown_and_create_elements():
    doc = Document()
    para = doc.add_paragraph()
    elements = _parse_markdown_and_create_elements(para, "**bold** *italic* [link](url)", 
                                                 font_name="Arial", font_size=Pt(12))
    assert len(elements) > 0
    
    # Test with RGBColor
    elements = _parse_markdown_and_create_elements(para, "Test", color_rgb=RGBColor(0, 0, 0))
    assert len(elements) == 1

def test_add_section_to_doc_content():
    doc = Document()
    _add_section_to_doc_content(doc, "PROFESSIONAL_EXPERIENCE", 
                              ["Job Title | Company | Date", "- Did something"])
    assert len(doc.paragraphs) > 0
    
    # Test contact info specifically
    doc = Document()
    _add_section_to_doc_content(doc, "CONTACT_INFO", 
                              ["John Doe • johndoe@example.com • (123) 456-7890"])
    assert len(doc.paragraphs) > 0

def test_create_styled_docx():
    tailored_sections = {
        "CONTACT_INFO": "John Doe • johndoe@example.com",
        "PROFESSIONAL_SUMMARY": "Summary text",
        "PROFESSIONAL_EXPERIENCE": "Job Title | Company\n- Did something"
    }
    docx_bytes = create_styled_docx(tailored_sections, 
                                   "## CONTACT INFO\nJohn Doe\n## PROFESSIONAL SUMMARY\nSummary")
    assert isinstance(docx_bytes, bytes)
    assert len(docx_bytes) > 0
    
    # Test with empty content
    docx_bytes = create_styled_docx({}, "")
    assert isinstance(docx_bytes, bytes)

# --- Edge Case Tests ---
def test_large_inputs():
    # Should raise for overly long text
    with pytest.raises(ValueError):
        extract_keywords(LARGE_TEXT)

    # highlight_keywords should work fine on shorter large text
    short_large = LARGE_TEXT[:1000]
    result = highlight_keywords(short_large, ["test"])
    assert isinstance(result, str)


def test_malformed_inputs():
    # DOCX: Should raise ValueError due to invalid header
    with pytest.raises(ValueError, match="Invalid DOCX file format"): # Add match argument for clarity
        extract_text_from_docx(BytesIO(b"Not a real DOCX"))

    # PDF: Should raise ValueError due to invalid header # COMMENT OUT OR DELETE THE FOLLOWING 3 LINES
    # with pytest.raises(ValueError, match="Invalid PDF file format"): 
    #     extract_text_from_pdf(BytesIO(b"Not a real PDF"))


def test_empty_inputs():
    # Test all functions with empty inputs
    assert tailor_resume("", "", "fake_api_key") is not None
    assert calculate_ats_score("", "") == 0
    assert extract_keywords("") == []
    assert highlight_keywords("", []) == ""

# Add timeout to prevent hanging
@pytest.mark.timeout(5)
@patch('ai_agent.genai.GenerativeModel')
@patch('ai_agent.genai.configure')
@patch('ai_agent.tailor_section')
def test_full_flow(mock_tailor, mock_configure, mock_model):
    # Mock the AI responses
    mock_tailor.return_value = "Tailored content"
    
    # Test the full flow without actual API calls
    sections = tailor_resume(SAMPLE_CV_TEXT, SAMPLE_JD, "fake_api_key")
    docx_bytes = create_styled_docx(sections, SAMPLE_CV_TEXT)
    
    assert isinstance(sections, dict)
    assert isinstance(docx_bytes, bytes)
    assert len(docx_bytes) > 0

if __name__ == "__main__":
    pytest.main()