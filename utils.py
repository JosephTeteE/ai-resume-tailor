# utils.py
"""
Utility functions for document generation and formatting.
Handles DOCX creation and file operations.
"""

import re
from io import BytesIO
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_TAB_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from config import MASTER_RESUME_DATA

def slugify(text):
    """
    Convert string to URL-friendly slug with title case.
    
    Args:
        text: Input string to convert
        
    Returns:
        str: Formatted slug
    """
    if not text:
        return ""
    
    text = str(text).strip()
    text = ' '.join(word.capitalize() for word in text.split())
    text = re.sub(r'[^\w\s-]', '', text)
    return re.sub(r'\s+', '_', text)

def _add_hyperlink(paragraph, text, url):
    """Add hyperlink to a paragraph while preserving formatting."""
    part = paragraph.part
    r_id = part.relate_to(url, RT.HYPERLINK, is_external=True)
    
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    
    # Preserve document font styling
    rFont = OxmlElement('w:rFonts')
    rFont.set(qn('w:ascii'), 'Times New Roman')
    rFont.set(qn('w:hAnsi'), 'Times New Roman')
    rPr.append(rFont)
    
    # Set font size and color
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), '20')
    rPr.append(sz)
    
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '0563C1')
    rPr.append(color)
    
    # Add underline
    underline = OxmlElement('w:u')
    underline.set(qn('w:val'), 'single')
    rPr.append(underline)
    
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)

def _add_section_header(doc, text):
    """Add formatted section header to document."""
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(10)
    run = para.add_run(text)
    run.bold = True
    run.underline = True
    return para

def create_final_docx(resume_data):
    """
    Generate professional resume DOCX file from structured data.
    
    Args:
        resume_data: Tailored resume content
        
    Returns:
        bytes: DOCX file in memory
    """
    doc = Document()
    
    # Configure base document styles
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(10)
    style.paragraph_format.line_spacing = 1.0
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(0)
    
    # Set page margins
    section = doc.sections[0]
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    
    # Configure right-aligned tab stop
    tab_stops = style.paragraph_format.tab_stops
    tab_stops.add_tab_stop(Inches(7.0), WD_TAB_ALIGNMENT.RIGHT)
    
    # Add contact information header
    contact = MASTER_RESUME_DATA['CONTACT_INFO']
    name_para = doc.add_paragraph()
    name_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    name_run = name_para.add_run(contact['name'])
    name_run.font.size = Pt(14)
    name_run.bold = True
    
    # Add contact details with email hyperlink
    details_para = doc.add_paragraph()
    details_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    details_para.paragraph_format.space_after = Pt(4)
    
    email_match = re.search(r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b', contact['details'])
    if email_match:
        details_para.add_run(contact['details'][:email_match.start()])
        _add_hyperlink(details_para, email_match.group(0), f"mailto:{email_match.group(0)}")
        details_para.add_run(contact['details'][email_match.end():])
    else:
        details_para.add_run(contact['details'])
    
    # Add divider line
    divider = doc.add_paragraph()
    divider_pr = divider._p.get_or_add_pPr()
    border = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'auto')
    border.append(bottom)
    divider_pr.append(border)
    
    # Add education section
    _add_section_header(doc, "EDUCATION")
    for edu in MASTER_RESUME_DATA['EDUCATION']:
        # Institution and location
        inst_para = doc.add_paragraph()
        inst_para.add_run(edu['institution']).bold = True
        inst_para.add_run(f"\t{edu['location']}")
        
        # Degree and dates
        deg_para = doc.add_paragraph()
        deg_run = deg_para.add_run(edu['degree'])
        deg_run.bold = True
        deg_run.italic = True
        deg_para.add_run(f"\t{edu['dates']}").bold = True
        
        # Relevant courses
        doc.add_paragraph(edu['courses']).add_run().italic = True
    
    # Add experience section
    _add_section_header(doc, "RELEVANT EXPERIENCE")
    for company, static_info in MASTER_RESUME_DATA["RELEVANT_EXPERIENCE_STATIC"].items():
        if company in resume_data.get('experience', {}):
            # Company header
            comp_para = doc.add_paragraph()
            comp_para.add_run(company).bold = True
            comp_para.add_run(f"\t{static_info['location']}")
            
            # Role and dates
            role_para = doc.add_paragraph()
            role_para.add_run(resume_data['experience'][company]['role']).bold = True
            role_para.add_run(f"\t{static_info['dates']}").bold = True
            
            # Bullet points
            for bullet in resume_data['experience'][company]['bullets']:
                doc.add_paragraph(bullet, style='List Bullet')
    
    # Add skills section
    _add_section_header(doc, "SKILLS")
    tech_para = doc.add_paragraph()
    tech_para.add_run("Technical Skills: ").bold = True
    tech_para.add_run(resume_data['skills']['technical'])
    
    soft_para = doc.add_paragraph()
    soft_para.paragraph_format.space_before = Pt(4)
    soft_para.add_run("Soft Skills: ").bold = True
    soft_para.add_run(resume_data['skills']['soft'])
    
    # Convert to bytes and return
    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio.getvalue()

def create_cover_letter_docx(text_content):
    """
    Generate simple DOCX file from cover letter text.
    
    Args:
        text_content: Plain text cover letter content
        
    Returns:
        bytes: DOCX file in memory
    """
    doc = Document()
    
    # Configure document styles
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(11)
    style.paragraph_format.line_spacing = 1.15
    
    # Set generous margins
    section = doc.sections[0]
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    
    # Add content preserving paragraphs
    for para in text_content.split('\n'):
        doc.add_paragraph(para)
    
    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio.getvalue()

def create_cheatsheet_docx(markdown_text):
    """
    Convert markdown-formatted text to formatted DOCX.
    
    Args:
        markdown_text: Cheatsheet content with markdown formatting
        
    Returns:
        bytes: DOCX file in memory
    """
    doc = Document()
    
    # Configure base style
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    
    # Parse markdown content
    for line in markdown_text.split('\n'):
        line = line.strip()
        if not line:
            continue
            
        # Handle headings
        if line.startswith('### '):
            doc.add_heading(line[4:].strip(), level=3)
        elif line.startswith('## '):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith('# '):
            doc.add_heading(line[2:].strip(), level=1)
        # Handle bullet points
        elif line.startswith(('• ', '* ', '- ')):
            doc.add_paragraph(line[2:].strip(), style='List Bullet')
        else:
            doc.add_paragraph(line)
    
    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio.getvalue()